from __future__ import annotations

from pathlib import Path

from django.conf import settings
from django.core.files import File
from django.utils import timezone
from docx import Document

from imports.models import FabImport
from reports.models import Report


def generate_word_report(fab_import: FabImport) -> dict:
    report_dir = Path(settings.MEDIA_ROOT) / "generated_reports" / str(fab_import.id)
    report_dir.mkdir(parents=True, exist_ok=True)

    clients_csv = report_dir / "clients_scores_code_relance_1.csv"
    zones_csv = report_dir / "zones_prioritaires_code_relance_1.csv"
    docx_path = report_dir / "rapport_recouvrement.docx"

    clients_df = __import__("pandas").DataFrame(list(fab_import.clients.order_by("-score_client").values()))
    zones_df = __import__("pandas").DataFrame(list(fab_import.zones.order_by("-score_zone").values()))
    clients_df.to_csv(clients_csv, index=False)
    zones_df.to_csv(zones_csv, index=False)

    document = Document()
    document.add_heading("Rapport quotidien du recouvrement", level=1)
    document.add_paragraph(f"Date: {timezone.localdate().strftime('%d/%m/%Y')}")
    document.add_paragraph(f"Fichier FAB: {fab_import.nom_fichier}")
    document.add_paragraph("Resume du traitement")
    document.add_paragraph(f"Lignes initiales: {fab_import.nombre_lignes_total}")
    document.add_paragraph(f"Clients retenus: {fab_import.nombre_clients_filtres}")
    document.add_paragraph(f"Zones generees: {fab_import.nombre_zones}")
    document.add_paragraph(f"Montant total: {fab_import.montant_total:,.2f} MRU")
    document.add_paragraph(f"Anciennete moyenne: {fab_import.anciennete_moyenne_jours:.2f} jours")
    document.add_paragraph(
        "Criteres de scoring: Client = 55% solde, 25% activite, 20% anciennete. "
        "Zone = 50% solde total, 25% nombre de clients, 15% anciennete, 10% activite."
    )
    document.add_paragraph("Top 10 zones prioritaires")
    for zone in fab_import.zones.order_by("-score_zone")[:10]:
        document.add_paragraph(
            f"{zone.zone_code} - score {zone.score_zone:.2f} - {zone.solde_total:,.2f} MRU",
            style="List Bullet",
        )

    repartition = {
        "HAUTE": fab_import.zones.filter(priorite_zone="HAUTE").count(),
        "MOYENNE": fab_import.zones.filter(priorite_zone="MOYENNE").count(),
        "FAIBLE": fab_import.zones.filter(priorite_zone="FAIBLE").count(),
    }
    document.add_paragraph("Repartition des priorites")
    for label, value in repartition.items():
        document.add_paragraph(f"{label}: {value}", style="List Bullet")

    document.add_paragraph("Recommandations terrain")
    document.add_paragraph("Prioriser les zones a score eleve.")
    document.add_paragraph("Planifier des campagnes ciblees sur les gros soldes.")
    document.add_paragraph("Conclusion")
    document.add_paragraph("Le dernier FAB met en evidence les zones et clients a traiter en priorite.")
    document.save(docx_path)

    with docx_path.open("rb") as generated_file:
        word_report = Report.objects.create(
            fab_import=fab_import,
            titre=f"Rapport quotidien du recouvrement - {fab_import.nom_fichier}",
            type_rapport=Report.TypeRapport.QUOTIDIEN,
            format=Report.FormatRapport.WORD,
            date_generation=timezone.now(),
        )
        word_report.fichier.save(docx_path.name, File(generated_file), save=True)

    with clients_csv.open("rb") as generated_file:
        clients_report = Report.objects.create(
            fab_import=fab_import,
            titre=f"Clients scores - {fab_import.nom_fichier}",
            type_rapport=Report.TypeRapport.CLIENTS,
            format=Report.FormatRapport.CSV,
            date_generation=timezone.now(),
        )
        clients_report.fichier.save(clients_csv.name, File(generated_file), save=True)

    with zones_csv.open("rb") as generated_file:
        zones_report = Report.objects.create(
            fab_import=fab_import,
            titre=f"Zones prioritaires - {fab_import.nom_fichier}",
            type_rapport=Report.TypeRapport.ZONES,
            format=Report.FormatRapport.CSV,
            date_generation=timezone.now(),
        )
        zones_report.fichier.save(zones_csv.name, File(generated_file), save=True)

    return {
        "word_report_id": word_report.id,
        "clients_report_id": clients_report.id,
        "zones_report_id": zones_report.id,
        "word_file": word_report.fichier.url,
        "clients_csv": clients_report.fichier.url,
        "zones_csv": zones_report.fichier.url,
    }
